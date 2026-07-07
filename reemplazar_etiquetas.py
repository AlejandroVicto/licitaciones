from docx import Document

def reemplazar_etiquetas(plantilla_path, salida_path, reemplazos):
    # Abrir el documento de Word
    doc = Document(plantilla_path)
    
    # Función auxiliar para reemplazar texto en párrafos
    def reemplazar_en_parrafos(parrafos):
        for p in parrafos:
            for etiqueta, valor in reemplazos.items():
                if etiqueta in p.text:
                    # NOTA: Asignar p.text directamente elimina el formato de ese párrafo. 
                    # Para conservar formato, se requeriría un manejo complejo de "runs".
                    p.text = p.text.replace(etiqueta, valor)

    # Buscar y reemplazar en los párrafos del documento principal
    reemplazar_en_parrafos(doc.paragraphs)
    
    # Buscar y reemplazar en las tablas (si las hay)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                reemplazar_en_parrafos(celda.paragraphs)
                
    # Guardar el documento modificado
    doc.save(salida_path)

if __name__ == "__main__":
    # Definir los reemplazos
    datos_a_reemplazar = {
        "{{REPRESENTANTE}}": "Juan Pérez",
        "{{RFC}}": "XAXX010101AAA"
    }
    
    # Archivos de entrada y salida
    archivo_entrada = "plantilla_prueba.docx"
    archivo_salida = "resultado_final.docx"
    
    print(f"Abriendo {archivo_entrada}...")
    reemplazar_etiquetas(archivo_entrada, archivo_salida, datos_a_reemplazar)
    print(f"Listo. Documento guardado como: {archivo_salida}")
