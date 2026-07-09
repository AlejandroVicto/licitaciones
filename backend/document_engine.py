import os
from docx import Document
#Reemplaza las etiquetas del documento con los datos de la empresa
def reemplazar_etiquetas(plantilla_path, salida_path, reemplazos):
    doc = Document(plantilla_path)
    
    def reemplazar_en_parrafos(parrafos):
        for p in parrafos:
            for etiqueta, valor in reemplazos.items():
                if etiqueta in p.text:
                    p.text = p.text.replace(etiqueta, valor)

    reemplazar_en_parrafos(doc.paragraphs)
    
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                reemplazar_en_parrafos(celda.paragraphs)
                
    doc.save(salida_path)

def generate_document(datos_empresa):
    archivo_plantilla = "plantilla_prueba.docx"
    archivo_salida = "resultado_final.docx"
    
    if not os.path.exists(archivo_plantilla):
        return None, f"No se encontró el archivo de plantilla: {archivo_plantilla}"
        
    reemplazos = {
        "{{REPRESENTANTE}}": datos_empresa.get('representante', ''),
        "{{RFC}}": datos_empresa.get('rfc', '')
    }
    
    try:
        reemplazar_etiquetas(archivo_plantilla, archivo_salida, reemplazos)
        with open(archivo_salida, "rb") as file:
            bytes_docx = file.read()
        nombre_sugerido = f"Escrito_{datos_empresa.get('nombre', 'Empresa').replace(' ', '_')}.docx"
        return bytes_docx, nombre_sugerido
    except Exception as e:
        return None, str(e)
