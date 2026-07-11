import os
import tempfile
from docx import Document

def generar_documento_word(empresa: dict, socios_text: str, domicilio_text: str) -> str:
    """
    Lee una plantilla de Word, reemplaza las etiquetas con los datos de la empresa,
    y guarda el resultado en un archivo temporal.
    Retorna la ruta absoluta del archivo generado.
    """
    # Ruta asumiendo que la carpeta templates está junto a este script
    base_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(base_dir, "templates", "machote_base.docx")
    
    # Si no existe, creamos un documento en blanco para evitar errores
    if not os.path.exists(template_path):
        os.makedirs(os.path.dirname(template_path), exist_ok=True)
        doc = Document()
        doc.add_paragraph("Plantilla autogenerada. Falta el archivo machote_base.docx original.")
        doc.add_paragraph("Razón Social: {{RAZON_SOCIAL}}")
        doc.add_paragraph("RFC: {{RFC}}")
        doc.add_paragraph("Domicilio: {{DOMICILIO}}")
        doc.add_paragraph("Socios:\n{{SOCIOS}}")
        doc.save(template_path)
    
    doc = Document(template_path)
    
    # Diccionario de reemplazos
    reemplazos = {
        "{{RAZON_SOCIAL}}": empresa.get("nombre", ""),
        "{{RFC}}": empresa.get("rfc", ""),
        "{{DOMICILIO}}": domicilio_text,
        "{{SOCIOS}}": socios_text,
        "{{PADRON}}": empresa.get("num_padron", ""),
        "{{EMAIL}}": empresa.get("email", ""),
        "{{TELEFONO}}": empresa.get("telefono", ""),
        "{{OBJETO_SOCIAL}}": empresa.get("objeto_social", ""),
        "{{ADMIN_UNICO}}": empresa.get("admin_unico", ""),
        "{{RFC_ADMIN}}": empresa.get("rfc_admin_unico", ""),
        "{{REPRESENTANTE}}": empresa.get("representante", ""),
        "{{RFC_REPRESENTANTE}}": empresa.get("rfc_representante", "")
    }
    
    # Asegurar que no haya valores nulos
    for k, v in reemplazos.items():
        if v is None:
            reemplazos[k] = ""
            
    # Reemplazar en párrafos
    for p in doc.paragraphs:
        for key, value in reemplazos.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))
                
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in reemplazos.items():
                        if key in p.text:
                            p.text = p.text.replace(key, str(value))
                            
    # Guardar en un archivo temporal
    temp_dir = tempfile.gettempdir()
    temp_file_path = os.path.join(temp_dir, f"documento_{empresa.get('id', 'temp')}.docx")
    doc.save(temp_file_path)
    
    return temp_file_path
